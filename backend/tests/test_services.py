"""
Unit tests for the file_processor service.
Uses only in-memory bytes — no disk I/O, no external services.
"""

import io
import pytest

from app.services.file_processor import extract_text, MAX_CHARS, SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# Plain text / Markdown
# ---------------------------------------------------------------------------

def test_extract_txt():
    """UTF-8 text bytes are returned as a string."""
    content = b"Hello, world! This is a plain text file."
    result = extract_text("notes.txt", content)
    assert result == "Hello, world! This is a plain text file."


def test_extract_txt_unicode():
    """UTF-8 characters (non-ASCII) are preserved correctly."""
    content = "Café résumé naïve".encode("utf-8")
    result = extract_text("unicode.txt", content)
    assert "Café" in result
    assert "résumé" in result


def test_extract_md():
    """Markdown content is extracted as raw text (no HTML conversion)."""
    content = b"# Heading\n\nSome **bold** text and a [link](http://example.com)."
    result = extract_text("readme.md", content)
    assert "Heading" in result
    assert "bold" in result


def test_extract_txt_invalid_bytes_replaced():
    """Invalid UTF-8 bytes are replaced rather than raising an error."""
    content = b"Valid text \xff\xfe invalid bytes"
    result = extract_text("bad.txt", content)
    # Should not raise; replacement characters present
    assert "Valid text" in result


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def test_extract_pdf():
    """A minimal valid PDF yields extracted text."""
    pypdf = pytest.importorskip("pypdf")
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    # Blank page has no text — just verify it doesn't raise
    result = extract_text("blank.pdf", pdf_bytes)
    assert isinstance(result, str)


def test_extract_pdf_with_text(tmp_path):
    """PDF with real text content is extracted correctly."""
    pypdf = pytest.importorskip("pypdf")
    from pypdf import PdfWriter
    from pypdf.generic import NameObject, DecodedStreamObject

    # Build a minimal single-page PDF with a text stream
    # We skip this complex low-level construction and use reportlab if available,
    # otherwise skip gracefully.
    reportlab = pytest.importorskip("reportlab.pdfgen.canvas", reason="reportlab not installed")
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Hello from PDF content")
    c.save()
    pdf_bytes = buf.getvalue()

    result = extract_text("test.pdf", pdf_bytes)
    assert "Hello from PDF content" in result


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def test_extract_docx():
    """A .docx file with paragraphs is extracted to text."""
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph in the document.")
    doc.add_paragraph("Second paragraph with more content.")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = extract_text("document.docx", docx_bytes)
    assert "First paragraph in the document." in result
    assert "Second paragraph with more content." in result


def test_extract_docx_empty():
    """An empty .docx (no paragraphs with text) returns an empty string."""
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = extract_text("empty.docx", docx_bytes)
    # python-docx adds a default empty paragraph; result should be empty after strip
    assert result == ""


# ---------------------------------------------------------------------------
# Unsupported extensions
# ---------------------------------------------------------------------------

def test_unsupported_extension_raises():
    """.exe files raise ValueError with a helpful message."""
    with pytest.raises(ValueError) as exc_info:
        extract_text("malware.exe", b"\x4d\x5a\x90")
    assert ".exe" in str(exc_info.value)
    assert "Unsupported" in str(exc_info.value)


def test_unsupported_extension_png():
    """.png files raise ValueError."""
    with pytest.raises(ValueError):
        extract_text("image.png", b"\x89PNG\r\n")


def test_unsupported_extension_no_extension():
    """Files with no extension raise ValueError."""
    with pytest.raises(ValueError):
        extract_text("Makefile", b"all:\n\techo done")


# ---------------------------------------------------------------------------
# Truncation
# ---------------------------------------------------------------------------

def test_truncation():
    """Content longer than MAX_CHARS is truncated with a note."""
    long_content = ("A" * (MAX_CHARS + 500)).encode("utf-8")
    result = extract_text("long.txt", long_content)

    assert len(result) > MAX_CHARS  # includes the truncation note
    assert f"[...truncated at {MAX_CHARS} chars]" in result
    # The actual content before truncation should be MAX_CHARS 'A' characters
    assert result.startswith("A" * MAX_CHARS)


def test_exact_max_chars_not_truncated():
    """Content at exactly MAX_CHARS is not truncated."""
    content = ("B" * MAX_CHARS).encode("utf-8")
    result = extract_text("exact.txt", content)
    assert result == "B" * MAX_CHARS
    assert "truncated" not in result


def test_short_content_not_truncated():
    """Content well under MAX_CHARS is returned as-is."""
    content = b"Short text."
    result = extract_text("short.txt", content)
    assert result == "Short text."
    assert "truncated" not in result


# ---------------------------------------------------------------------------
# Empty file
# ---------------------------------------------------------------------------

def test_empty_file_txt():
    """Empty bytes for a .txt file returns an empty string."""
    result = extract_text("empty.txt", b"")
    assert result == ""


def test_whitespace_only_file():
    """Whitespace-only content is stripped to an empty string."""
    result = extract_text("spaces.txt", b"   \n\n\t  \n")
    assert result == ""


# ---------------------------------------------------------------------------
# Supported extensions set
# ---------------------------------------------------------------------------

def test_supported_extensions_set():
    """The SUPPORTED_EXTENSIONS constant contains all expected types."""
    assert ".txt" in SUPPORTED_EXTENSIONS
    assert ".md" in SUPPORTED_EXTENSIONS
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS
